from pathlib import Path
from io import BytesIO

from pypdf import PdfReader
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ResumeParserError(Exception):
    """Raised when a resume cannot be parsed."""
    pass


def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract text from a PDF stored in memory.
    """

    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception as exc:
        raise ResumeParserError(
            f"Unable to read PDF file: {exc}"
        ) from exc

    pages = []

    for page in reader.pages:
        try:
            text = page.extract_text() or ""
            pages.append(text)
        except Exception:
            pages.append("")

    extracted_text = "\n".join(pages).strip()

    if not extracted_text:
        raise ResumeParserError(
            "No readable text was found in the PDF. "
            "The resume may be scanned/image-based."
        )

    return extracted_text


def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file stored in memory.
    """

    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise ResumeParserError(
            f"Unable to read DOCX file: {exc}"
        ) from exc

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    # Also extract text from tables.
    table_text = []

    for table in document.tables:
        for row in table.rows:
            row_values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if row_values:
                table_text.append(" | ".join(row_values))

    all_text = paragraphs + table_text

    extracted_text = "\n".join(all_text).strip()

    if not extracted_text:
        raise ResumeParserError(
            "No readable text was found in the DOCX file."
        )

    return extracted_text


def extract_resume_text(
    filename: str,
    file_bytes: bytes
) -> str:
    """
    Detect file type and extract resume text.
    """

    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ResumeParserError(
            f"Unsupported file type: {extension}. "
            f"Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if extension == ".pdf":
        return extract_pdf_text(file_bytes)

    if extension == ".docx":
        return extract_docx_text(file_bytes)

    raise ResumeParserError("Unsupported resume format.")